import asyncio
from playwright.async_api import async_playwright
import openpyxl
from bs4 import BeautifulSoup
import re
from docx import Document
import os
import time
import json

new_folder_path = 'C:\\Users\\phams\\Downloads\\doanh_nghiep\\thong tin doanh nghiep'
os.makedirs(new_folder_path, exist_ok=True)

types = [
    ('chu-dau-tu',10),
    ('nha-thau-tu-van',2),
    ('tong-thau-du-an',4),
    ('thau-co-dien-vat-tu-thi-cong',1),
    ('thau-hang-muc-xay-dung-vat-tu-thi-cong',1),
    ('dich-vu-logicstic',1),
    ('cong-nghe-va-thiet-bi-san-xuat',1)
]

# types = [('gian-hang',2)]

for obj in types:
    spreadsheet = openpyxl.load_workbook(f'C:\\Users\\phams\\Downloads\\doanh_nghiep\\{obj[0]}.xlsx')
    sheet = spreadsheet.active
    
    def extract_info(idx):
        with open(f'C:\\Users\\phams\\Downloads\\doanh_nghiep\\page_sources\\{obj[0]}_{idx-1}.html', 'r', encoding='utf-8') as file:
            html_content = file.read()

        # Parse the HTML content
        soup = BeautifulSoup(html_content, 'html.parser')

        checkpoint_text = None
        detail_article = soup.find('div', class_='d-flex d-md-none row')
        if detail_article:
            h1_tag = detail_article.find('h1')
            if h1_tag:
                checkpoint_text = h1_tag.get_text(separator='\n').strip()

        if checkpoint_text:
            checkpoint_text = str(checkpoint_text).replace('/','-')
            
            # Create a new Document
            doc = Document()
            
            doc.add_heading(checkpoint_text,level=1)
            
            # Find the div with class 'special-class'
            special_div = soup.find('div', class_='col-12 col-md-8')

            # Extract all text inside that div
            if special_div:
                for text_element in special_div.stripped_strings:
                    doc.add_paragraph(text_element)
                    
            # Find the div with class 'special-class'
            special_div = soup.find('div', class_='col-12 col-md-4 mt-3')

            # Extract all text inside that div
            if special_div:
                for text_element in special_div.stripped_strings:
                    doc.add_paragraph(text_element)

            # Save the document
            output_file_path = f'C:\\Users\\phams\\Downloads\\doanh_nghiep\\thong tin doanh nghiep\\{str(checkpoint_text)}_{idx-1}.docx'
            doc.save(output_file_path)
            
            sheet.cell(row=idx,column=3).value = checkpoint_text

            print(f"Texts extracted, cleaned, and saved successfully to {output_file_path}.")
        else:
            print(f"Checkpoint text not found: {sheet.cell(row=idx,column=2).value}")

    for i in range(2,sheet.max_row+1):
        extract_info(i)

    sheet_file = f"C:\\Users\\phams\\Downloads\\doanh_nghiep\\{obj[0]}.xlsx"
    spreadsheet.save(sheet_file)