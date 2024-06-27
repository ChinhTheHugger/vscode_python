import requests
from shapely.geometry import LineString, Point
from shapely.ops import split
from shapely import wkt
import numpy as np
import json
import polyline
import ast
import matplotlib.pyplot as plt
import contextily as ctx
import geopandas as gpd
import psycopg2
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from openpyxl import Workbook
from openpyxl.drawing.image import Image
import io
import os
import pandas as pd
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from PIL import Image
from pyproj import Proj, transform
from geojson_length import calculate_distance, Unit
from geojson import Feature, LineString as GeoJSONLineString
import openpyxl
from scipy.spatial.distance import directed_hausdorff
from docx.shared import Pt
from scipy.spatial.distance import euclidean
from fastdtw import fastdtw
import networkx as nx
from scipy.spatial.distance import cdist
import time



# Google excel sheet
spreadsheet_gg = openpyxl.load_workbook("C:\\Users\\phams\\Downloads\\failed_case.xlsx")

sheet_gg = spreadsheet_gg.active

# result document
doc = Document()

# result sheet
spreadsheet = openpyxl.Workbook()
sheet = spreadsheet.active

img_stream = io.BytesIO()



def graphhopper_api(start,end,type):
    url = "https://graphhopper.com/api/1/route"

    query = {
    "key": "2bb8da46-a001-441b-815f-83d023d4d912"
    }

    payload = {
    "profile": type,
    "points": [
        [
            start[1], start[0]
        ],
        [
            end[1], end[0]
        ]
    ], # long, lat format
    "points_encoded": True, # False returns a non encoded string of coordinates
    "optimize": "true",
    "point_hints": [
        "Lindenschmitstraße",
        "Thalkirchener Str."
    ],
    "snap_preventions": [
        "motorway",
        "ferry",
        "tunnel"
    ],
    "details": [
        "road_class",
        "surface"
    ]
    }

    headers = {"Content-Type": "application/json"}

    response = requests.post(url, json=payload, headers=headers, params=query)

    data = response.json()
    
    if data.get('paths') is None:
        data = {
            'status': False,
            'route': '',
            'message': data['message']
        }
        
        return data
    else:
        encoded_points = data['paths'][0]['points']
        decoded_points = polyline.decode(encoded_points)
        
        data = {
            'status': True,
            'route': [(lon, lat) for lat, lon in decoded_points],
            'message': "No error"
        }
        
        return data

# Define a function to simplify a LineString
def simplify_route(line, tolerance=0.001):
    return line.simplify(tolerance, preserve_topology=False)

# Define a function to compute the Hausdorff distance
def hausdorff_distance(line1, line2):
    u = np.array(line1.coords)
    v = np.array(line2.coords)
    return max(directed_hausdorff(u, v)[0], directed_hausdorff(v, u)[0])

# Define a function to compute the average Hausdorff distance
def average_hausdorff_distance(line1, line2):
    u = np.array(line1.coords)
    v = np.array(line2.coords)
    distances_uv = [min(np.linalg.norm(u_i - v, axis=1)) for u_i in u]
    distances_vu = [min(np.linalg.norm(v_i - u, axis=1)) for v_i in v]
    avg_distance = (sum(distances_uv) + sum(distances_vu)) / (len(distances_uv) + len(distances_vu))
    return avg_distance

# Define a function to convert degrees to meters (for latitude and longitude)
def degrees_to_meters(degrees, latitude=0):
    meters_per_degree_lat = 111e3  # Approximate conversion: 1 degree of latitude = 111 km
    meters_per_degree_lon = meters_per_degree_lat * np.cos(np.radians(latitude))
    return degrees * meters_per_degree_lat, degrees * meters_per_degree_lon

def convert_to_lon_lat(linestring):
    return [[lon, lat] for lon, lat in linestring.coords]

def split_route_into_sections(line, percentage):
    total_length = line.length
    section_length = total_length * percentage / 100.0
    num_sections = int(np.ceil(total_length / section_length))
    points = [line.interpolate(float(n) / num_sections, normalized=True) for n in range(num_sections + 1)]
    sections = [(points[i], points[i+1]) for i in range(len(points) - 1)]
    return sections

def check_section_overlap(route_section, route):
    section_line = LineString(route_section)
    return section_line.intersects(route)

def check_route_through_sections(route_1, route_2, percentage):
    sections_1 = split_route_into_sections(route_1, percentage)
    
    for section in sections_1:
        if check_section_overlap(section, route_2):
            return True
    
    return False

# Haversine distance function
def haversine(lonlat1, lonlat2):
    R = 6371.0  # Earth radius in kilometers
    
    lon1, lat1 = np.radians(lonlat1[:, 0]), np.radians(lonlat1[:, 1])
    lon2, lat2 = np.radians(lonlat2[:, 0]), np.radians(lonlat2[:, 1])
    
    dlon = lon2 - lon1[:, np.newaxis]
    dlat = lat2 - lat1[:, np.newaxis]
    
    a = np.sin(dlat / 2)**2 + np.cos(lat1[:, np.newaxis]) * np.cos(lat2) * np.sin(dlon / 2)**2
    c = 2 * np.arcsin(np.sqrt(a))
    
    return R * c

# Define a function to interpolate points along a LineString
def interpolate_points(line, num_points):
    distances = np.linspace(0, line.length, num_points)
    return np.array([line.interpolate(distance).coords[0] for distance in distances])



doc.add_heading('LineString Analysis', level=1)

for i in range(2,20):
    print("Case: "+str(sheet_gg.cell(row=i,column=1).value))
    doc.add_heading(f"Case: {str(sheet_gg.cell(row=i,column=1).value)}", level=2)

    start = (sheet_gg.cell(row=i,column=3).value,sheet_gg.cell(row=i,column=2).value) # lat, long
    end = (sheet_gg.cell(row=i,column=5).value,sheet_gg.cell(row=i,column=4).value) # lat, long
    
    doc.add_paragraph(f"Start at: lat = {sheet_gg.cell(row=i,column=3).value}, long = {sheet_gg.cell(row=i,column=2).value}")
    doc.add_paragraph(f"End at: lat = {sheet_gg.cell(row=i,column=5).value}, long = {sheet_gg.cell(row=i,column=4).value}")
    
    num_points = 250  # Number of points to interpolate

    # GraphHopper API accepts: car, bike, foot
    gh_response_car = graphhopper_api(start,end,'car')
    time.sleep(10)
    gh_response_bike = graphhopper_api(start,end,'bike')
    time.sleep(10)
    gh_response_foot = graphhopper_api(start,end,'foot')
    time.sleep(10)
    
    if gh_response_car['status'] is False:
        gh_line_car = LineString([])
        error_message_car = gh_response_car['message']
    else:
        gh_line_car = LineString(gh_response_car['route'])
        error_message_car = gh_response_car['message']
    
    if gh_response_bike['status'] is False:
        gh_line_bike = LineString([])
        error_message_bike = gh_response_bike['message']
    else:
        gh_line_bike = LineString(gh_response_bike['route'])
        error_message_bike = gh_response_bike['message']
    
    if gh_response_foot['status'] is False:
        gh_line_foot = LineString([])
        error_message_foot = gh_response_foot['message']
    else:
        gh_line_foot = LineString(gh_response_foot['route'])
        error_message_foot = gh_response_foot['message']
    
    google_line = wkt.loads(sheet_gg.cell(row=i,column=7).value)
    
    interpolated_b = interpolate_points(google_line, num_points)
    
    ls_data = {
        'geometry': [google_line,gh_line_car,gh_line_bike,gh_line_foot],
        'route_name': ['Google','GraphHopper car','GraphHopper bike','GraphHopper foot'],
        'error_message': ['No error',error_message_car,error_message_bike,error_message_foot],
        'colors': ['red','blue','green','orange']
    }
    
    dis_data = {
        'LineStrings': {},
        'Reference': {'interpolated_points': interpolated_b}
    }
    
    dis_color = ['blue','green','orange']
    
    ls_gdf = gpd.GeoDataFrame(data=ls_data, crs="EPSG:4326")
    
    for x in range(1,len(ls_gdf)):
        ls = ls_gdf.iloc[x].geometry
        name = ls_gdf.iloc[x].route_name
        error = ls_gdf.iloc[x].error_message
        
        doc.add_heading(f'{name} summary:', level=3)
        
        doc.add_heading('Route Lengths', level=4)
        
        if not LineString(ls).is_empty:
            ls_converted = convert_to_lon_lat(ls)
            line = Feature(geometry=GeoJSONLineString(ls_converted))
            length = calculate_distance(line,Unit.kilometers)
            
            print(f"{name}: {error}")
            
            doc.add_paragraph(f"Length of {name} route: {length:.2f} kilometer")
            sheet.cell(row=i,column=6+x).value = length
        else:
            print(f"{name}: {error}")
            
            doc.add_paragraph(f"{name} didn't return a value: {error}")
            
        doc.add_heading('Route Distances', level=4)
        
        if not LineString(ls).is_empty:
            # Interpolate additional points along each LineString
            
            interpolated_a = interpolate_points(ls, num_points)
            
            # Calculate pairwise haversine distances for interpolated points
            distances_ab = haversine(interpolated_a, interpolated_b)

            # Find the minimum distance for each point in interpolated_a to any point in interpolated_b
            min_distances_a_to_b = distances_ab.min(axis=1)

            # Calculate key statistics
            max_distance = np.max(min_distances_a_to_b)
            min_distance = np.min(min_distances_a_to_b)
            average_distance = np.mean(min_distances_a_to_b)

            # Thresholds in kilometers
            major_divergence_threshold = 0.015 # meters

            # Identify significant peaks
            major_divergences = np.where(min_distances_a_to_b > major_divergence_threshold)[0]
            
            doc.add_paragraph(f'Maximum distance: {max_distance:.2f} km')
            doc.add_paragraph(f'Minimum distance: {min_distance:.2f} km')
            doc.add_paragraph(f'Average distance: {average_distance:.2f} km')
            
            if len(major_divergences) > 0:
                doc.add_paragraph(f'Percentage of sections of {name} with major divergence from Google line (distance > 15 meters, which is bigger than the width of a 4 lanes road): {len(major_divergences) * 100 / num_points} %')
            
            dis_data['LineStrings'][name] = {
                'interpolated_points': interpolated_a,
                'distances_to_b': min_distances_a_to_b,
                'color': dis_color[x-1]
            }
        else:
            doc.add_paragraph(f"{name} didn't return a value, so distance to Google route can't be calculated")
        
    # Create a figure with 2 subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 6))

    # Plot the distance graph on the first subplot
    for name, values in dis_data['LineStrings'].items():
        ax1.plot(range(num_points), values['distances_to_b'], marker='o', linestyle='-', label=f'{name} to B', color=values['color'], markersize=1, linewidth=1)

    ax1.axhline(y=major_divergence_threshold, color='r', linestyle='--', label='Major Divergence Threshold', linewidth=1)
    ax1.set_title('Shortest Distance from Each Point on LineStrings A to LineString B')
    ax1.set_xlabel('Index of Point on LineString A')
    ax1.set_ylabel('Minimum Distance to LineString B (km)')
    ax1.legend()
    ax1.grid(True)
    
    for idx, row in ls_gdf.iterrows():
        # Plot the LineString
        if not row['geometry'].is_empty:
            ls_gdf.loc[[idx]].plot(ax=ax2, label=row['route_name'], color=row['colors'], linewidth=2)

    # Customize and show the plot
    ax2.set_title('Route from Point A to Point B')
    ax2.set_xlabel('Longitude')
    ax2.set_ylabel('Latitude')
    ax2.legend(loc='lower right')
    # ctx.add_basemap(ax2, crs="EPSG:4326", source=ctx.providers.OpenStreetMap.Mapnik, zoom=15)
    ax2.grid(True)

    # manager = plt.get_current_fig_manager()
    # manager.window.state('zoomed')

    # Adjust layout
    plt.tight_layout()
    # plt.show()

    plt.savefig(img_stream, format='png')
    # Save the plot to a file
    plot_file = "C:\\Users\\phams\\Downloads\\linestrings_plot.png"
    fig.savefig(plot_file)
    plt.close(fig)  # Close the plot

    # Check if the file is created and valid
    # try:
    #     with Image.open(plot_file) as img:
    #         img.show()  # Open the image to verify
    # except IOError:
    #     print("Error: The image was not saved correctly.")

    doc.add_heading('Plot of routes and distances', level=2)
    doc.add_picture(plot_file, width=Inches(5))
    
    # doc.add_page_break()
    
# Define the desired font size
font_size = Pt(12)  # 12 point font size
# Set the font size for all paragraphs and runs in the document
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        run.font.size = font_size
doc_file = "C:\\Users\\phams\\Downloads\\linestring_analysis_test.docx"
doc.save(doc_file)