import asyncio
from playwright.async_api import async_playwright
import openpyxl
from bs4 import BeautifulSoup
import re
from docx import Document

spreadsheet = openpyxl.load_workbook("C:\\Users\\phams\\Downloads\\Bac_Ninh\\bac_ninh_projects_links.xlsx")
sheet = spreadsheet.active

def extract_info(idx):
    with open(f'C:\\Users\\phams\\Downloads\\Bac_Ninh\\page sources\\page_source_{idx-1}.html', 'r', encoding='utf-8') as file:
        html_content = file.read()

    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Extract the checkpoint_text from the <h1> tag within the <article class="detail"> tag
    checkpoint_text = None
    detail_article = soup.find('article', class_='detail')
    if detail_article:
        h1_tag = detail_article.find('h1')
        if h1_tag:
            checkpoint_text = h1_tag.get_text(separator='\n').strip()

    if checkpoint_text:
        # print(f"Checkpoint text found: {checkpoint_text}")

        # Extract text from <article> and <tbody> tags
        article_texts = [tag.get_text(separator='\n') for tag in soup.find_all('article')]
        tbody_texts = [tag.get_text(separator='\n') for tag in soup.find_all('tbody')]

        # Combine the texts
        all_texts = article_texts + tbody_texts

        # Remove duplicates
        unique_texts = list(set(all_texts))

        # Join texts with a single line break and clean up
        cleaned_texts = [re.sub(r'\n+', '\n', text).strip() for text in unique_texts]

        # Remove leading and trailing spaces from each line in the cleaned texts
        cleaned_texts = [re.sub(r'^\s+', '', text, flags=re.MULTILINE) for text in cleaned_texts]

        # Ensure checkpoint_text is the first entry
        if checkpoint_text in cleaned_texts:
            index = cleaned_texts.index(checkpoint_text)
            cleaned_texts = cleaned_texts[index:]

        # Create a new Document
        doc = Document()

        # Add cleaned texts to the document
        for text in cleaned_texts:
            doc.add_paragraph(text)
            doc.add_paragraph('\n')

        # Save the document
        output_file_path = f'C:\\Users\\phams\\Downloads\\Bac_Ninh\\thong tin cac du an\\{checkpoint_text}.docx'
        doc.save(output_file_path)
        
        sheet.cell(row=idx,column=3).value = checkpoint_text

        print(f"Texts extracted, cleaned, and saved successfully to {output_file_path}.")
    else:
        print("Checkpoint text not found.")

# async def get_accessibility_tree(url, idx):
#     async with async_playwright() as p:
#         browser = await p.firefox.launch_persistent_context(user_data_dir="C:\\Users\\phams\\AppData\\Roaming\\Mozilla\\Firefox\\Profiles\\bd5r7ma0.default-release - Copy", headless=True)
#         page = await browser.new_page()
#         await page.goto(url)
        
#         # # Get the accessibility tree
#         # accessibility_snapshot = await page.accessibility.snapshot()
        
#         # with open('C:\\Users\\phams\\Downloads\\accessibility_tree.json', 'w') as f:
#         #     json.dump(accessibility_snapshot, f, indent=2)
        
#         # await browser.close()
        
#         # Get the page source
#         page_source = await page.content()
        
#         source_file_path = f'C:\\Users\\phams\\Downloads\\Bac_Ninh\\page sources\\page_source_{idx-1}.html'
        
#         # Save the page source to an HTML file
#         with open(source_file_path, 'w', encoding='utf-8') as f:
#             f.write(page_source)
            
#         try:
#             extract_info(source_file_path,idx)
#         except FileNotFoundError:
#             print(f"File '{source_file_path}' not found.")
#         except Exception as e:
#             print(f"Error reading file: {e}")
        
#         await browser.close()

for i in range(2,sheet.max_row+1):
    extract_info(i)

sheet_file = "C:\\Users\\phams\\Downloads\\Bac_Ninh\\bac_ninh_projects_links.xlsx"
spreadsheet.save(sheet_file)
