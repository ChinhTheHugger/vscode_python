import requests
from shapely.geometry import LineString, Point
from shapely.ops import split
from shapely import wkt
import numpy as np
import json
import polyline
import ast
import matplotlib.pyplot as plt
import contextily as ctx
import geopandas as gpd
import psycopg2
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from openpyxl import Workbook
from openpyxl.drawing.image import Image
import io
import os
import pandas as pd
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from PIL import Image
from pyproj import Proj, transform
from geojson_length import calculate_distance, Unit
from geojson import Feature, LineString as GeoJSONLineString
import openpyxl



# GHTK database connection
conn = psycopg2.connect(database="ccdb",
                        host="10.110.69.95",
                        user="gmap_user",
                        password="v78xzTArbexG8KKbmGaLvpTe",
                        port="5432")
cursor = conn.cursor()

cursor.execute(
    '''select * from gg_routing_logs'''
)

# Google sheet connection
scope = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"
]

creds_path = "C:\\Users\\phams\\Downloads\\mythic-evening-425602-k5-bcaf0ca4be0f.json"

creds = Credentials.from_service_account_file(creds_path, scopes=scope)

client = gspread.authorize(creds)

# excel connection
spreadsheet = openpyxl.load_workbook("C:\\Users\\phams\\Downloads\\failed_case.xlsx")

sheet = spreadsheet.active

doc = Document()

img_stream = io.BytesIO()



# GHTK API, accept: car, bike, motorcycle, xteam_motorcycle
def ghtk_gh_api(start,end,type):
    cookies = {
        'Phpstorm-1bdbdc0b': 'e072668f-73ea-4ab2-b595-4517c460b36e',
        '_osm_location': '105.86035|20.99552|19|M',
        'iconSize': '32x32',
        'jenkins-timestamper-offset': '-25200000',
    }

    headers = {
        'Connection': 'keep-alive',
        'sec-ch-ua': '" Not A;Brand";v="99", "Chromium";v="96", "Google Chrome";v="96"',
        'Accept': 'application/json, text/javascript, */*; q=0.01',
        'GH-Client': 'web-ui 3.0',
        'sec-ch-ua-mobile': '?0',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.93 Safari/537.36',
        'sec-ch-ua-platform': '"Linux"',
        'Sec-Fetch-Site': 'same-origin',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': 'http://localhost:8989/maps/?point=15.23119%2C108.262024&point=14.756291%2C108.492737&locale=en-US&elevation=false&profile=car&use_miles=false&layer=OpenStreetMap',
        'Accept-Language': 'en-US,en;q=0.9',
        # 'Cookie': 'Phpstorm-1bdbdc0b=e072668f-73ea-4ab2-b595-4517c460b36e; _osm_location=105.86035|20.99552|19|M; iconSize=32x32; jenkins-timestamper-offset=-25200000',
        'apikey': 'CgmUZhzdammE3A2guUgUXSyj',
        'Authorization': 'Bearer eyJhbGciOiJFUzI1NiIsImtpZCI6IjAxRjVOMThESE02RlkwSEpKSFhFRlE5NzNLXzE2MjA5ODIzODAiLCJ0eXAiOiJKV1QifQ.eyJhdWQiOiJhdXRoIiwiZXhwIjoxNjU2NjkxMjQ1LCJqdGkiOiIwMUc2WDRFTjlFOUM4RDNZSFRKQVlFQjM3NiIsImlhdCI6MTY1NjY4NzY0NSwiaXNzIjoiaHR0cHM6Ly9hdXRoLmdodGtsYWIuY29tIiwic3ViIjoiMDFHMlhNWTQzN1MxNEIyWVY0WDQzNlRTUkIiLCJzY3AiOlsib3BlbmlkIl0sInNpZCI6IlE1VFJFZFpXaHJ6SFJEa3JyRVE1bkM1dzRaRU93aDV3IiwiY2xpZW50X2lkIjoiMDFGNU4xOERITTZGWTBISkpIWEVGUTk3M0siLCJ0eXBlIjoib2F1dGgifQ.rq7Jj0zEBSq3nlEs5AZ1BAz5UL6BYbrjDz7QMyAnrXdVHPAqU9pfZ-VMXTbTrWfyxCB2h1pohcoBKbY4tOYvnQ',
        'Content-Type': 'application/json',
    }

    json_data = {
        'gh_requests': [
            {
                'points': [
                    [
                        start[1], start[0]
                    ],
                    [
                        end[1], end[0]
                    ],
                ],
                'vehicle': type, # car, bike, motorcycle, xteam_motorcycle
                'request_id': '982884285',
                'calc_points': True,
                'points_encoded': True,
                'instructions': True,
                'locale': 'vi',
                'algorithm': 'alternative_route',
            },
        ],
    }

    response = requests.post('https://gmap-api-gw.ghtklab.com/route', cookies=cookies, headers=headers, json=json_data)
    # print(json.dumps(response.json(),indent=4))

    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{\n    "gh_requests": [\n        {\n            "points": [\n                [\n                    105.863657,\n                    20.98368\n                ],\n                [\n                    105.861919,\n                    20.997323\n                ]\n            ],\n            "vehicle": "motorcycle",\n            "request_id": "982884285",\n            "calc_points": true,\n            "points_encoded": true,\n            "instructions": true,\n            "locale": "vi",\n            "algorithm": "alternative_route"\n        }\n    ]\n}'
    #response = requests.post('https://gmap-api-gw.ghtklab.com/route', cookies=cookies, headers=headers, data=data)

    data = response.json()
    if data.get('gh_responses') is None:
        return False
    else:
        gh_response = data['gh_responses']
        if gh_response[0].get('paths', None) is None or gh_response[0]['paths'] == []:
            return False
        else:
            paths_data = gh_response[0]['paths']
            encoded_points = paths_data[0]['points']
            decoded_points = polyline.decode(encoded_points)
            
            return [(lon, lat) for lat, lon in decoded_points]

# GraphHopper API, accept: car, bike, foot
def osm_gh_api(start,end,type):
    url = "https://graphhopper.com/api/1/route"

    query = {
    "key": "669ddb50-80a8-4bc3-94ed-6371ca6dc8c7"
    }

    payload = {
    "profile": type,
    "points": [
        [
            start[1], start[0]
        ],
        [
            end[1], end[0]
        ]
    ], # long, lat format
    "points_encoded": True, # False returns a non encoded string of coordinates
    "optimize": "true",
    "point_hints": [
        "Lindenschmitstraße",
        "Thalkirchener Str."
    ],
    "snap_preventions": [
        "motorway",
        "ferry",
        "tunnel"
    ],
    "details": [
        "road_class",
        "surface"
    ]
    }

    headers = {"Content-Type": "application/json"}

    response = requests.post(url, json=payload, headers=headers, params=query)

    data = response.json()
    if data.get('paths') is None:
        return False
    else:
        paths_data = data['paths']
        if paths_data[0].get('points', None) is None or paths_data[0]['points'] == []:
            return False
        else:
            encoded_points = paths_data[0]['points']
            decoded_points = polyline.decode(encoded_points)
            
            return [(lon, lat) for lat, lon in decoded_points]

# Function to sample points along a LineString
def sample_points(linestring, num_points):
    distances = np.linspace(0, linestring.length, num_points)
    points = [linestring.interpolate(distance) for distance in distances]
    return points

# Function to calculate the average distance between points on two LineStrings
def calculate_average_distance(ls1, ls2, num_points=100):
    points1 = sample_points(ls1, num_points)
    points2 = sample_points(ls2, num_points)
    
    distances = [point1.distance(point2) for point1, point2 in zip(points1, points2)]
    average_distance = np.mean(distances)
    return average_distance

def convert_to_lon_lat(linestring):
    return [[lon, lat] for lon, lat in linestring.coords]

def split_route_into_sections(line, percentage):
    total_length = line.length
    section_length = total_length * percentage / 100.0
    num_sections = int(np.ceil(total_length / section_length))
    points = [line.interpolate(float(n) / num_sections, normalized=True) for n in range(num_sections + 1)]
    sections = [(points[i], points[i+1]) for i in range(len(points) - 1)]
    return sections

def check_section_overlap(route_section, route):
    section_line = LineString(route_section)
    return section_line.intersects(route)

def check_route_through_sections(route_1, route_2, percentage):
    sections_1 = split_route_into_sections(route_1, percentage)
    
    for section in sections_1:
        if check_section_overlap(section, route_2):
            return True
    
    return False



doc.add_heading('LineString Analysis', level=1)

for i in range(2,sheet.max_row+1):
    print("Case: "+str(sheet.cell(row=i,column=1).value))
    doc.add_heading(f"Case: {str(sheet.cell(row=i,column=1).value)}", level=2)

    start = (sheet.cell(row=i,column=3).value,sheet.cell(row=i,column=2).value) # lat, long
    end = (sheet.cell(row=i,column=5).value,sheet.cell(row=i,column=4).value) # lat, long
    
    doc.add_paragraph(f"Start at: lat = {sheet.cell(row=i,column=3).value}, long = {sheet.cell(row=i,column=2).value}")
    doc.add_paragraph(f"End at: lat = {sheet.cell(row=i,column=5).value}, long = {sheet.cell(row=i,column=4).value}")

    # constant
    osm_data_c = osm_gh_api(start,end,'car')
    osm_data_b = osm_gh_api(start,end,'bike')
    osm_data_f = osm_gh_api(start,end,'foot')

    if osm_data_c is not False:
        osm_line_c = LineString(osm_data_c)
    else:
        osm_line_c = LineString()

    if osm_data_b is not False:
        osm_line_b = LineString(osm_data_b)
    else:
        osm_line_b = LineString()

    if osm_data_f is not False:
        osm_line_f = LineString(osm_data_f)
    else:
        osm_line_f = LineString()
    
    google_line = wkt.loads(sheet.cell(row=i,column=7).value)
    
    data = {
        'geometry': [osm_line_c,osm_line_b,osm_line_f,google_line],
        'route_name': ['OSM car','OSM bike','OSM foot','Google']
    }
    
    gdf = gpd.GeoDataFrame(data=data, crs="EPSG:4326")
    
    colors = ['blue','green','purple','red']

    fig, ax = plt.subplots()

    for idx, row in gdf.iterrows():
        # Plot the LineString
        if not row['geometry'].is_empty:
            gdf.loc[[idx]].plot(ax=ax, label=row['route_name'], color=colors[idx % len(colors)], linewidth=2)

    # Set a fixed aspect ratio
    ax.set_aspect('equal')

    # Customize and show the plot
    plt.title('Route from Point A to Point B')
    plt.xlabel('Longitude')
    plt.ylabel('Latitude')
    plt.legend(loc='lower right')
    ctx.add_basemap(ax, crs="EPSG:4326", source=ctx.providers.OpenStreetMap.Mapnik, zoom=15)
    plt.grid(True)

    doc.add_heading('Route Lengths', level=3)

    gdf_projected = gdf.to_crs("EPSG:3857")

    for i in range(len(gdf)):
        ls = gdf.iloc[i].geometry
        name = gdf.iloc[i].route_name
        
        if not LineString(ls).is_empty:
            ls_converted = convert_to_lon_lat(ls)
            line = Feature(geometry=GeoJSONLineString(ls_converted))
            length = calculate_distance(line,Unit.kilometers)
            
            print(f"Length of {name} route: {length:.2f} kilometer")
            
            doc.add_paragraph(f"Length of {name} route: {length:.2f} kilometer")
        else:
            print(f"{name} didn't return a value")
            
            doc.add_paragraph(f"{name} didn't return a value")

    doc.add_heading('Average Distances', level=3)

    for i in range(len(gdf_projected)):
        ls1 = gdf_projected.iloc[i].geometry
        name1 = gdf_projected.iloc[i].route_name
        
        if LineString(ls1).is_empty:
            print(f"{name1} didn't return a value, therefore average distance to other routes can't be calculated")
            
            doc.add_paragraph(f"{name1} didn't return a value, therefore average distance to other routes can't be calculated")
        else:
            for j in range(i + 1, len(gdf_projected)):
                ls2 = gdf_projected.iloc[j].geometry
                name2 = gdf_projected.iloc[j].route_name
                
                if LineString(ls2).is_empty:
                    print(f"{name2} didn't return a value, therefore average distance to {name1} can't be calculated")
            
                    doc.add_paragraph(f"{name2} didn't return a value, therefore average distance to {name1} can't be calculated")
                else:
                    avg_distance = calculate_average_distance(ls1, ls2)
                    overlap = check_route_through_sections(ls1,ls2,5.0)
                    
                    print(f"Average distance between {name1} and {name2}: {avg_distance:.2f} meters")
                    
                    doc.add_paragraph(f"Average distance between {name1} and {name2}: {avg_distance:.2f} meters")
                    
                    if overlap:
                        print(f"The {name1} route goes through one or more sections of the {name2} route")
                        
                        doc.add_paragraph(f"The {name1} route goes through one or more sections of the {name2} route")
                    else:
                        print(f"The {name1} route does not go through any sections of the {name2} route")
                        
                        doc.add_paragraph(f"The {name1} route does not go through any sections of the {name2} route")

    # manager = plt.get_current_fig_manager()
    # manager.window.state('zoomed')

    # plt.show()

    plt.savefig(img_stream, format='png')
    # Save the plot to a file
    plot_file = "C:\\Users\\phams\\Downloads\\linestrings_plot.png"
    fig.savefig(plot_file)
    plt.close(fig)  # Close the plot

    # Check if the file is created and valid
    # try:
    #     with Image.open(plot_file) as img:
    #         img.show()  # Open the image to verify
    # except IOError:
    #     print("Error: The image was not saved correctly.")

    doc.add_heading('Plot of LineStrings', level=2)
    doc.add_picture(plot_file, width=Inches(5))
    
    # doc.add_page_break()

doc_file = "C:\\Users\\phams\\Downloads\\linestring_analysis.docx"
doc.save(doc_file)

