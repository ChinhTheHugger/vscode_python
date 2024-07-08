import openpyxl
import polyline
from shapely import LineString
import numpy as np
import math
import requests
from docx import Document
import io
import matplotlib.pyplot as plt
import geopandas as gpd
from docx.shared import Inches



spreadsheet_gg = openpyxl.load_workbook("C:\\Users\\phams\\Downloads\\failed_case.xlsx")
sheet_gg = spreadsheet_gg.active

doc = Document()
doc_graph = Document()

img_stream = io.BytesIO()



# Haversine distance function - return in kilometer
def haversine(lon1, lat1, lon2, lat2):
    # Radius of the Earth in km
    R = 6371.0

    # Convert decimal degrees to radians
    lat1 = math.radians(lat1)
    lon1 = math.radians(lon1)
    lat2 = math.radians(lat2)
    lon2 = math.radians(lon2)

    # Differences in coordinates
    dlat = lat2 - lat1
    dlon = lon2 - lon1

    # Haversine formula
    a = math.sin(dlat / 2)**2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2)**2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    # Distance in kilometers
    distance = R * c

    return distance

def ghtk_gh_api_local(start,end,type):
    headers = {
        'apikey': 'aghzzzzzzzzzzzzui',
        'Content-Type': 'application/json',
        'Authorization': 'Bearer eyJhbGciOiJFUzI1NiIsImtpZCI6IjAxRjRTMFBTRjdQR001VDZDWVExTTVNRVg3XzE2MjAwNDIyNzgiLCJ0eXAiOiJKV1QifQ.eyJhdWQiOiJnbWFwIiwiZXhwIjoxNjcwMzg4OTg0LCJqdGkiOiIwMUdLSldOS1NTN0FDSldTTVExMURHNzhZRSIsImlhdCI6MTY3MDMwMjU3OSwiaXNzIjoiaHR0cHM6Ly9hdXRoLmdodGtsYWIuY29tIiwic3ViIjoiMDFGNFMwUFNGN1BHTTVUNkNZUTFNNU1FWDciLCJzY3AiOlsiZ21hcDphZGRyZXNzLXZlcmlmaWNhdGlvbi1hcGkuYWNjZXNzLWRpcmVjdCIsImdtYXA6YWRkcmVzc2VzLW1hbmFnZW1lbnQuYWNjZXNzLWRpcmVjdCIsImdtYXA6Y2FwYWNpdHkuYWNjZXNzLWRpcmVjdCIsImdtYXA6Z29vZ2xlLW1hcC1odHRwLXNlcnZpY2UuYWNjZXNzLWRpcmVjdCIsImdtYXA6aGFtbGV0LW5vZGUtZGV0YWlsLmFjY2Vzcy1kaXJlY3QiLCJnbWFwOmludGVybmFsLWFwaS5hY2Nlc3MtZGlyZWN0IiwiZ21hcDpwcmVnZW8uYWNjZXNzLWRpcmVjdCIsImdtYXA6cm91dGUuYWNjZXNzLWRpcmVjdCIsImdtYXA6c3VnZ2VzdC1hZGRyZXNzLmFjY2Vzcy1kaXJlY3QiXSwiY2xpZW50X2lkIjoiMDFGNFMwUFNGN1BHTTVUNkNZUTFNNU1FWDciLCJ0eXBlIjoiZGlyZWN0In0.b89nX-mK6pjtfAMtqcCXFghvPC1qbOHjPDaDDrz9ljuTpdAhkvxkVlnST2tSyxNJ4dluwjHIBPsmFzWS5erJUA',
    }

    json_data = {
        'gh_requests': [
            {
                'points': [
                    [
                        start[1],start[0]
                    ],
                    [
                        end[1],end[0]
                    ],
                ],
                'profile': type,
                'request_id': '982884285',
                'calc_points': True,
                'points_encoded': True,
                'instructions': True,
            },
        ],
    }

    response = requests.post('http://10.110.69.186:8989/route', headers=headers, json=json_data)

    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{\n        "gh_requests": [\n            {\n                "points": [\n                    [\n                        105.842638,\n                        21.200526\n                    ],\n                    [\n                        104.1166,\n                        21.185161\n                    ]\n                ],\n                "profile": "car",\n                "request_id": "982884285",\n                "calc_points": true,\n                "points_encoded": true,\n                "instructions": true\n            }\n        ]\n    }'
    #response = requests.post('http://10.110.69.186:8989/route', headers=headers, data=data)

    data = response.json()
    
    if data.get('gh_responses') is None:
        data = {
            'status': False,
            'route': '',
            'message': "Error getting response"
        }
        
        return data
    else:
        gh_response = data['gh_responses']
        if gh_response[0].get('paths', None) is None or gh_response[0]['paths'] == []:
            data = {
                'status': False,
                'route': '',
                'message': gh_response[0]['errors'][0]['message']
            }
            
            return data
        else:
            paths_data = gh_response[0]['paths']
            encoded_points = paths_data[0]['points']
            decoded_points = polyline.decode(encoded_points)
            
            data = {
                'status': True,
                'route': [(lon, lat) for lat, lon in decoded_points],
                'message': "No error"
            }
            
            return data



cases = [1,2,13,18,20,22,24,25,26,27,28,30,32,33,34,35,37,38,41,43,47,49,51,54,55,56,58,59,61,62,66,68,69,70,71,72,74]

for i in range(2,sheet_gg.max_row+1):
    if int(sheet_gg.cell(row=i,column=1).value) in cases:
        doc_graph.add_heading(f"Case: {str(sheet_gg.cell(row=i,column=1).value)}", level=2)
        
        start = (sheet_gg.cell(row=i,column=3).value,sheet_gg.cell(row=i,column=2).value) # lat, long
        end = (sheet_gg.cell(row=i,column=5).value,sheet_gg.cell(row=i,column=4).value) # lat, long
        
        google_encoded = str(sheet_gg.cell(row=i,column=6).value)
        google_decoded = polyline.decode(google_encoded)
        google_line = LineString(google_decoded)
        google_line = LineString([[lon, lat] for lat, lon in google_line.coords])
        
        gh_response_foot = ghtk_gh_api_local(start,end,'motorcycle')
        
        if gh_response_foot['status'] is False:
            gh_line_foot = LineString([])
            error_message_foot = gh_response_foot['message']
        else:
            gh_line_foot = LineString(gh_response_foot['route'])
            error_message_foot = gh_response_foot['message']
        
        ls_data = {
            'geometry': [google_line,gh_line_foot],
            'route_name': ['Google','GH motor'],
            'error_message': ['No error',error_message_foot],
            'colors': ['red','blue']
        }
        
        fig, (ax2) = plt.subplots(1, 1, figsize=(15, 15))

        
        ls_gdf = gpd.GeoDataFrame(data=ls_data, crs="EPSG:4326")
        
        for idx, row in ls_gdf.iterrows():
            # Plot the LineString
            if row['route_name'] == 'GH motor' or row['route_name'] == 'Google':
                if not row['geometry'].is_empty:
                    ls_gdf.loc[[idx]].plot(ax=ax2, label=row['route_name'], color=row['colors'], linewidth=1)

        # x1, y1 = google_line.coords[0][1], google_line.coords[0][0]
        # x2, y2 = google_line.coords[-1][1], google_line.coords[-1][0]
            
        # plt.plot(x1, y1, marker='o', markersize=5, color='red')
        # plt.plot(x2, y2, marker='o', markersize=5, color='red')
        
        # x1, y1 = gh_line_foot.coords[0][1], gh_line_foot.coords[0][0]
        # x2, y2 = gh_line_foot.coords[-1][1], gh_line_foot.coords[-1][0]
            
        # plt.plot(x1, y1, marker='o', markersize=5, color='blue')
        # plt.plot(x2, y2, marker='o', markersize=5, color='blue')
        
        x1, y1 = start[1], start[0]
        x2, y2 = end[1], end[0]
        
        plt.plot(x1, y1, marker='o', markersize=5, color='green')
        plt.plot(x2, y2, marker='o', markersize=5, color='green')

        # Customize and show the plot
        ax2.set_title('Route from Point A to Point B')
        ax2.set_xlabel('Longitude')
        ax2.set_ylabel('Latitude')
        ax2.legend(loc='lower right')
        # ctx.add_basemap(ax2, crs="EPSG:4326", source=ctx.providers.OpenStreetMap.Mapnik, zoom=15)
        ax2.grid(True)
        
        plt.savefig(img_stream, format='png')
        # Save the plot to a file
        plot_file = "C:\\Users\\phams\\Downloads\\linestrings_plot_check.png"
        fig.savefig(plot_file)
        plt.close(fig)  # Close the plot
        
        doc_graph.add_picture(plot_file, width=Inches(5))

doc_file_graph = "C:\\Users\\phams\\Downloads\\linestring_analysis_graphs.docx"
doc_graph.save(doc_file_graph)