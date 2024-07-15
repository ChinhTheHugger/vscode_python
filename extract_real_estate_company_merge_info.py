import asyncio
from playwright.async_api import async_playwright
import openpyxl
from bs4 import BeautifulSoup
import re
from docx import Document
import os
import time
import json

new_folder_path = 'C:\\Users\\phams\\Downloads\\doanh_nghiep\\thong tin doanh nghiep\\tong hop'
os.makedirs(new_folder_path, exist_ok=True)

types = [
    ('nha-thau-tu-van',2),
    ('tong-thau-du-an',4),
    ('thau-co-dien-vat-tu-thi-cong',1),
    ('thau-hang-muc-xay-dung-vat-tu-thi-cong',1),
    ('dich-vu-logicstic',1),
    ('cong-nghe-va-thiet-bi-san-xuat',1),
    ('chu-dau-tu',10)
]

# types = [('gian-hang',2)]

for obj in types:
    spreadsheet = openpyxl.load_workbook(f'C:\\Users\\phams\\Downloads\\doanh_nghiep\\{obj[0]}.xlsx')
    sheet = spreadsheet.active
    
    merge_doc = Document()
    
    check_doc = []
    
    for i in range(2,sheet.max_row+1):
        if sheet.cell(row=i,column=3).value is None:
            continue
        else:
            source_doc = Document(f'C:\\Users\\phams\\Downloads\\doanh_nghiep\\thong tin doanh nghiep\\{sheet.cell(row=i,column=3).value}_{sheet.cell(row=i,column=1).value}.docx')
            
            if sheet.cell(row=i,column=3).value not in check_doc:
                merge_doc.add_heading(f'{sheet.cell(row=i,column=1).value - 1} - {sheet.cell(row=i,column=3).value}',level=1)
                
                for paragraph in source_doc.paragraphs:
                    merge_doc.add_paragraph(paragraph.text)
                    
                check_doc.append(sheet.cell(row=i,column=3).value)
            else:
                check_doc.clear()
                
                continue
            
            merge_doc.add_paragraph('')
        
    merge_doc.save(f'C:\\Users\\phams\\Downloads\\doanh_nghiep\\thong tin doanh nghiep\\tong hop\\{obj[0]}.docx')
    
    print(f"Merged info for {obj[0]}")