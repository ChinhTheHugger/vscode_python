import asyncio
from playwright.async_api import async_playwright
import openpyxl
from bs4 import BeautifulSoup
import re
from docx import Document
import os
import time

json_list = [('ha_nam',2)]

for pair in json_list:
    name, count = pair[0], pair[1]
    
    parts = name.split('_')
    capitalized_parts = [part.capitalize() for part in parts]
    
    capitalized_name = '_'.join(capitalized_parts)
    capitalized_name_with_space = ' '.join(capitalized_parts)
    
    spreadsheet = openpyxl.load_workbook(f"C:\\Users\\phams\\Downloads\\{capitalized_name}\\{name}_projects_links.xlsx")
    sheet = spreadsheet.active

    merge_doc = Document()

    for i in range(2,sheet.max_row+1):
        position = sheet.cell(row=i,column=1).value
        file_name = sheet.cell(row=i,column=3).value
        
        merge_doc.add_heading(f'{position} - {file_name}',level=1)
        
        source_doc = Document(f'C:\\Users\\phams\\Downloads\\{capitalized_name}\\thong tin du an\\{file_name}.docx')
        
        for paragraph in source_doc.paragraphs:
            merge_doc.add_paragraph(paragraph.text)
        
        merge_doc.add_paragraph('')

    merge_doc.save(f"C:\\Users\\phams\\Downloads\\{capitalized_name}\\{capitalized_name_with_space}.docx")
    
    print(f"Merged info for {capitalized_name_with_space}")