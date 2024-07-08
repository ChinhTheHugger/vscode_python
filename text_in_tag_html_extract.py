from bs4 import BeautifulSoup
import re
from docx import Document

# Load the HTML file
file_path = 'C:\\Users\\phams\\Downloads\\FCIVIETNAM - Thông tin thị trường, dự án Xây dựng - Nhà máy sản xuất.htm'
with open(file_path, 'r', encoding='utf-8') as file:
    html_content = file.read()

# Parse the HTML content
soup = BeautifulSoup(html_content, 'html.parser')

# Extract the checkpoint_text from the <h1> tag within the <article class="detail"> tag
checkpoint_text = None
detail_article = soup.find('article', class_='detail')
if detail_article:
    h1_tag = detail_article.find('h1')
    if h1_tag:
        checkpoint_text = h1_tag.get_text(separator='\n').strip()

if checkpoint_text:
    print(f"Checkpoint text found: {checkpoint_text}")

    # Extract text from <article> and <tbody> tags
    article_texts = [tag.get_text(separator='\n') for tag in soup.find_all('article')]
    tbody_texts = [tag.get_text(separator='\n') for tag in soup.find_all('tbody')]

    # Combine the texts
    all_texts = article_texts + tbody_texts

    # Remove duplicates
    unique_texts = list(set(all_texts))

    # Join texts with a single line break and clean up
    cleaned_texts = [re.sub(r'\n+', '\n', text).strip() for text in unique_texts]

    # Remove leading and trailing spaces from each line in the cleaned texts
    cleaned_texts = [re.sub(r'^\s+', '', text, flags=re.MULTILINE) for text in cleaned_texts]

    # Ensure checkpoint_text is the first entry
    if checkpoint_text in cleaned_texts:
        index = cleaned_texts.index(checkpoint_text)
        cleaned_texts = cleaned_texts[index:]

    # Create a new Document
    doc = Document()

    # Add cleaned texts to the document
    doc.add_paragraph(cleaned_texts[0])
    doc.add_paragraph('\n')

    # Save the document
    output_file_path = f'C:\\Users\\phams\\Downloads\\{checkpoint_text}.docx'
    doc.save(output_file_path)

    print(f"Texts extracted, cleaned, and saved successfully to {output_file_path}.")
else:
    print("Checkpoint text not found.")
