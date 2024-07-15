import requests
from shapely.geometry import LineString, Point, mapping, shape
from shapely.ops import split, transform, unary_union
from shapely import wkt
import numpy as np
import json
import polyline
import ast
import matplotlib.pyplot as plt
import contextily as ctx
import geopandas as gpd
import psycopg2
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from openpyxl import Workbook
from openpyxl.drawing.image import Image
import io
import os
import pandas as pd
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from PIL import Image
from pyproj import Proj, Transformer
from geojson_length import calculate_distance, Unit
from geojson import Feature, LineString as GeoJSONLineString
import openpyxl
from scipy.spatial.distance import directed_hausdorff
from docx.shared import Pt
from scipy.spatial.distance import euclidean
from fastdtw import fastdtw
import networkx as nx
from scipy.spatial.distance import cdist
import time
from scipy.optimize import fsolve
import math
import pyproj
from functools import partial
from datetime import date



# Google excel sheet
spreadsheet_gg = openpyxl.load_workbook("C:\\Users\\phams\\Downloads\\failed_case.xlsx")

sheet_gg = spreadsheet_gg.active

# result document
doc = Document()
doc_graph = Document()

# result sheet
spreadsheet = openpyxl.Workbook()
sheet = spreadsheet.active

img_stream = io.BytesIO()

today = date.today()



def graphhopper_api(start,end,type):
    url = "https://graphhopper.com/api/1/route"

    query = {
    "key": "2bb8da46-a001-441b-815f-83d023d4d912"
    }

    payload = {
    "profile": type, # car, bike, foot
    "points": [
        [
            start[1], start[0]
        ],
        [
            end[1], end[0]
        ]
    ], # long, lat format
    "points_encoded": True, # False returns a non encoded string of coordinates
    "optimize": "true",
    "point_hints": [
        "Lindenschmitstraße",
        "Thalkirchener Str."
    ],
    "snap_preventions": [
        "motorway",
        "ferry",
        "tunnel"
    ],
    "details": [
        "road_class",
        "surface"
    ]
    }

    headers = {"Content-Type": "application/json"}

    response = requests.post(url, json=payload, headers=headers, params=query)

    data = response.json()
    
    if data.get('paths') is None:
        data = {
            'status': False,
            'route': '',
            'message': data['message']
        }
        
        return data
    else:
        encoded_points = data['paths'][0]['points']
        decoded_points = polyline.decode(encoded_points)
        
        data = {
            'status': True,
            'route': [(lon, lat) for lat, lon in decoded_points],
            'message': "No error"
        }
        
        return data
    
def openrouteservice_api(start,end,type):
    headers = {
        'Content-Type': 'application/json; charset=utf-8',
        'Accept': 'application/json, application/geo+json, application/gpx+xml, img/png; charset=utf-8',
        'Authorization': '5b3ce3597851110001cf624882d0e6ed8c5c4aa28b9c89a1157d53f6',
    }

    json_data = {
        'coordinates': [
            [
                start[1], start[0]
            ],
            [
                end[1], end[0]
            ]
        ],
    }
    # type: driving-car, driving-hgv, cycling-regular, cycling-road, cycling-mountain, cycling-electric, foot-walking, foot-hiking, wheelchair
    response = requests.post(f'https://api.openrouteservice.org/v2/directions/{type}/json', headers=headers, json=json_data)

    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{"coordinates":[[8.681495,49.41461],[8.687872,49.420318]]}'
    #response = requests.post('https://api.openrouteservice.org/v2/directions/driving-car/json', headers=headers, data=data)
    
    data = response.json()
    
    if data.get('routes') is None:
        data = {
            'status': False,
            'route': '',
            'message': data['error']['message']
        }
        
        return data
    else:
        encoded_points = data['routes'][0]['geometry']
        decoded_points = polyline.decode(encoded_points)
        
        data = {
            'status': True,
            'route': [(lon, lat) for lat, lon in decoded_points],
            'message': "No error"
        }
        
        return data

def ghtk_gh_api(start,end,type):
    cookies = {
        'Phpstorm-1bdbdc0b': 'e072668f-73ea-4ab2-b595-4517c460b36e',
        '_osm_location': '105.86035|20.99552|19|M',
        'iconSize': '32x32',
        'jenkins-timestamper-offset': '-25200000',
    }

    headers = {
        'Connection': 'keep-alive',
        'sec-ch-ua': '" Not A;Brand";v="99", "Chromium";v="96", "Google Chrome";v="96"',
        'Accept': 'application/json, text/javascript, */*; q=0.01',
        'GH-Client': 'web-ui 3.0',
        'sec-ch-ua-mobile': '?0',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.93 Safari/537.36',
        'sec-ch-ua-platform': '"Linux"',
        'Sec-Fetch-Site': 'same-origin',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Dest': 'empty',
        'Referer': 'http://localhost:8989/maps/?point=15.23119%2C108.262024&point=14.756291%2C108.492737&locale=en-US&elevation=false&profile=car&use_miles=false&layer=OpenStreetMap',
        'Accept-Language': 'en-US,en;q=0.9',
        # 'Cookie': 'Phpstorm-1bdbdc0b=e072668f-73ea-4ab2-b595-4517c460b36e; _osm_location=105.86035|20.99552|19|M; iconSize=32x32; jenkins-timestamper-offset=-25200000',
        'apikey': 'CgmUZhzdammE3A2guUgUXSyj',
        'Authorization': 'Bearer eyJhbGciOiJFUzI1NiIsImtpZCI6IjAxRjVOMThESE02RlkwSEpKSFhFRlE5NzNLXzE2MjA5ODIzODAiLCJ0eXAiOiJKV1QifQ.eyJhdWQiOiJhdXRoIiwiZXhwIjoxNjU2NjkxMjQ1LCJqdGkiOiIwMUc2WDRFTjlFOUM4RDNZSFRKQVlFQjM3NiIsImlhdCI6MTY1NjY4NzY0NSwiaXNzIjoiaHR0cHM6Ly9hdXRoLmdodGtsYWIuY29tIiwic3ViIjoiMDFHMlhNWTQzN1MxNEIyWVY0WDQzNlRTUkIiLCJzY3AiOlsib3BlbmlkIl0sInNpZCI6IlE1VFJFZFpXaHJ6SFJEa3JyRVE1bkM1dzRaRU93aDV3IiwiY2xpZW50X2lkIjoiMDFGNU4xOERITTZGWTBISkpIWEVGUTk3M0siLCJ0eXBlIjoib2F1dGgifQ.rq7Jj0zEBSq3nlEs5AZ1BAz5UL6BYbrjDz7QMyAnrXdVHPAqU9pfZ-VMXTbTrWfyxCB2h1pohcoBKbY4tOYvnQ',
        'Content-Type': 'application/json',
    }

    json_data = {
        'gh_requests': [
            {
                'points': [
                    [
                        start[0],start[1]
                    ],
                    [
                        end[0],end[1]
                    ],
                ],
                'vehicle': type, # car, bike, motorcycle, xteam_motorcycle
                'request_id': '982884285',
                'calc_points': True,
                'points_encoded': True,
                'instructions': True,
                'locale': 'vi',
                'algorithm': 'alternative_route',
            },
        ],
    }

    response = requests.post('https://gmap-api-gw.ghtklab.com/route', cookies=cookies, headers=headers, json=json_data)
    # print(json.dumps(response.json(),indent=4))

    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{\n    "gh_requests": [\n        {\n            "points": [\n                [\n                    105.863657,\n                    20.98368\n                ],\n                [\n                    105.861919,\n                    20.997323\n                ]\n            ],\n            "vehicle": "motorcycle",\n            "request_id": "982884285",\n            "calc_points": true,\n            "points_encoded": true,\n            "instructions": true,\n            "locale": "vi",\n            "algorithm": "alternative_route"\n        }\n    ]\n}'
    #response = requests.post('https://gmap-api-gw.ghtklab.com/route', cookies=cookies, headers=headers, data=data)

    data = response.json()
    
    if data.get('gh_responses') is None:
        data = {
            'status': False,
            'route': '',
            'message': "Error getting response"
        }
        
        return data
    else:
        gh_response = data['gh_responses']
        if gh_response[0].get('paths', None) is None or gh_response[0]['paths'] == []:
            data = {
                'status': False,
                'route': '',
                'message': gh_response[0]['errors'][0]['message']
            }
            
            return data
        else:
            paths_data = gh_response[0]['paths']
            encoded_points = paths_data[0]['points']
            decoded_points = polyline.decode(encoded_points)
            
            data = {
                'status': True,
                'route': [(lon, lat) for lat, lon in decoded_points],
                'message': "No error"
            }
            
            return data

def ghtk_gh_api_local(start,end,type):
    headers = {
        'apikey': 'aghzzzzzzzzzzzzui',
        'Content-Type': 'application/json',
        'Authorization': 'Bearer eyJhbGciOiJFUzI1NiIsImtpZCI6IjAxRjRTMFBTRjdQR001VDZDWVExTTVNRVg3XzE2MjAwNDIyNzgiLCJ0eXAiOiJKV1QifQ.eyJhdWQiOiJnbWFwIiwiZXhwIjoxNjcwMzg4OTg0LCJqdGkiOiIwMUdLSldOS1NTN0FDSldTTVExMURHNzhZRSIsImlhdCI6MTY3MDMwMjU3OSwiaXNzIjoiaHR0cHM6Ly9hdXRoLmdodGtsYWIuY29tIiwic3ViIjoiMDFGNFMwUFNGN1BHTTVUNkNZUTFNNU1FWDciLCJzY3AiOlsiZ21hcDphZGRyZXNzLXZlcmlmaWNhdGlvbi1hcGkuYWNjZXNzLWRpcmVjdCIsImdtYXA6YWRkcmVzc2VzLW1hbmFnZW1lbnQuYWNjZXNzLWRpcmVjdCIsImdtYXA6Y2FwYWNpdHkuYWNjZXNzLWRpcmVjdCIsImdtYXA6Z29vZ2xlLW1hcC1odHRwLXNlcnZpY2UuYWNjZXNzLWRpcmVjdCIsImdtYXA6aGFtbGV0LW5vZGUtZGV0YWlsLmFjY2Vzcy1kaXJlY3QiLCJnbWFwOmludGVybmFsLWFwaS5hY2Nlc3MtZGlyZWN0IiwiZ21hcDpwcmVnZW8uYWNjZXNzLWRpcmVjdCIsImdtYXA6cm91dGUuYWNjZXNzLWRpcmVjdCIsImdtYXA6c3VnZ2VzdC1hZGRyZXNzLmFjY2Vzcy1kaXJlY3QiXSwiY2xpZW50X2lkIjoiMDFGNFMwUFNGN1BHTTVUNkNZUTFNNU1FWDciLCJ0eXBlIjoiZGlyZWN0In0.b89nX-mK6pjtfAMtqcCXFghvPC1qbOHjPDaDDrz9ljuTpdAhkvxkVlnST2tSyxNJ4dluwjHIBPsmFzWS5erJUA',
    }

    json_data = {
        'gh_requests': [
            {
                'points': [
                    [
                        start[1],start[0]
                    ],
                    [
                        end[1],end[0]
                    ],
                ],
                'profile': type,
                'request_id': '982884285',
                'calc_points': True,
                'points_encoded': True,
                'instructions': True,
            },
        ],
    }

    response = requests.post('http://10.110.69.186:8989/route', headers=headers, json=json_data)

    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{\n        "gh_requests": [\n            {\n                "points": [\n                    [\n                        105.842638,\n                        21.200526\n                    ],\n                    [\n                        104.1166,\n                        21.185161\n                    ]\n                ],\n                "profile": "car",\n                "request_id": "982884285",\n                "calc_points": true,\n                "points_encoded": true,\n                "instructions": true\n            }\n        ]\n    }'
    #response = requests.post('http://10.110.69.186:8989/route', headers=headers, data=data)

    data = response.json()
    
    if data.get('gh_responses') is None:
        data = {
            'status': False,
            'route': '',
            'message': "Error getting response"
        }
        
        return data
    else:
        gh_response = data['gh_responses']
        if gh_response[0].get('paths', None) is None or gh_response[0]['paths'] == []:
            data = {
                'status': False,
                'route': '',
                'message': gh_response[0]['errors'][0]['message']
            }
            
            return data
        else:
            paths_data = gh_response[0]['paths']
            encoded_points = paths_data[0]['points']
            decoded_points = polyline.decode(encoded_points)
            
            data = {
                'status': True,
                'route': [(lon, lat) for lat, lon in decoded_points],
                'message': "No error"
            }
            
            return data

# Define a function to simplify a LineString
def simplify_route(line, tolerance=0.001):
    return line.simplify(tolerance, preserve_topology=False)

# Define a function to compute the Hausdorff distance
def hausdorff_distance(line1, line2):
    u = np.array(line1.coords)
    v = np.array(line2.coords)
    return max(directed_hausdorff(u, v)[0], directed_hausdorff(v, u)[0])

# Define a function to compute the average Hausdorff distance
def average_hausdorff_distance(line1, line2):
    u = np.array(line1.coords)
    v = np.array(line2.coords)
    distances_uv = [min(np.linalg.norm(u_i - v, axis=1)) for u_i in u]
    distances_vu = [min(np.linalg.norm(v_i - u, axis=1)) for v_i in v]
    avg_distance = (sum(distances_uv) + sum(distances_vu)) / (len(distances_uv) + len(distances_vu))
    return avg_distance

# Define a function to convert degrees to meters (for latitude and longitude)
def degrees_to_meters(degrees, latitude=0):
    meters_per_degree_lat = 111e3  # Approximate conversion: 1 degree of latitude = 111 km
    meters_per_degree_lon = meters_per_degree_lat * np.cos(np.radians(latitude))
    return degrees * meters_per_degree_lat, degrees * meters_per_degree_lon

def convert_to_lon_lat(linestring):
    return [[lon, lat] for lon, lat in linestring.coords]

def split_route_into_sections(line, percentage):
    total_length = line.length
    section_length = total_length * percentage / 100.0
    num_sections = int(np.ceil(total_length / section_length))
    points = [line.interpolate(float(n) / num_sections, normalized=True) for n in range(num_sections + 1)]
    sections = [(points[i], points[i+1]) for i in range(len(points) - 1)]
    return sections

def check_section_overlap(route_section, route):
    section_line = LineString(route_section)
    return section_line.intersects(route)

def check_route_through_sections(route_1, route_2, percentage):
    sections_1 = split_route_into_sections(route_1, percentage)
    
    for section in sections_1:
        if check_section_overlap(section, route_2):
            return True
    
    return False

# Haversine distance function
def haversine(lonlat1, lonlat2):
    R = 6371.0  # Earth radius in kilometers
    
    lon1, lat1 = np.radians(lonlat1[:, 0]), np.radians(lonlat1[:, 1])
    lon2, lat2 = np.radians(lonlat2[:, 0]), np.radians(lonlat2[:, 1])
    
    dlon = lon2 - lon1[:, np.newaxis]
    dlat = lat2 - lat1[:, np.newaxis]
    
    a = np.sin(dlat / 2)**2 + np.cos(lat1[:, np.newaxis]) * np.cos(lat2) * np.sin(dlon / 2)**2
    c = 2 * np.arcsin(np.sqrt(a))
    
    return R * c

# Haversine distance function - return in kilometer
def haversine_for_points(point1,point2):
    # Radius of the Earth in km
    R = 6371.0

    # Convert decimal degrees to radians
    lat1 = math.radians(point1[1])
    lon1 = math.radians(point1[0])
    lat2 = math.radians(point2[1])
    lon2 = math.radians(point2[0])

    # Differences in coordinates
    dlat = lat2 - lat1
    dlon = lon2 - lon1

    # Haversine formula
    a = math.sin(dlat / 2)**2 + math.cos(lat1) * math.cos(lat2) * math.sin(dlon / 2)**2
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    # Distance in kilometers
    distance = R * c

    return distance

# Define a function to interpolate points along a LineString
def interpolate_points(line, num_points):
    distances = np.linspace(0, line.length, num_points)
    return np.array([line.interpolate(distance).coords[0] for distance in distances])

def overlay_analysis(line1, line2):
    intersection = line1.intersection(line2)
    union = line1.union(line2)
    difference = line1.difference(line2)
    
    result = {
        'intersection': intersection,
        'union': union,
        'difference': difference
    }
    
    return result

# Function to convert degrees to radians
def degrees_to_radians(degrees):
    return degrees * math.pi / 180.0

# Define EPSG:4326 and a suitable projected coordinate system, e.g., UTM
proj_wgs84 = Proj(init='epsg:4326')
proj_utm = Proj(proj='utm', zone=18, ellps='WGS84')

# Define transformer for EPSG:4326 to UTM
transformer_to_utm = Transformer.from_proj(proj_wgs84, proj_utm)
transformer_to_wgs84 = Transformer.from_proj(proj_utm, proj_wgs84)

# Function to transform coordinates
def to_utm(geom):
    return transform(transformer_to_utm.transform, geom)

def to_wgs84(geom):
    return transform(transformer_to_wgs84.transform, geom)



doc.add_heading('LineString Analysis', level=1)
doc_graph.add_heading('LineString Analysis', level=1)

sheet.cell(row=1,column=1).value = 'Case'

sheet.cell(row=1,column=2).value = 'Start lat'
sheet.cell(row=1,column=3).value = 'Start long'

sheet.cell(row=1,column=4).value = 'End lat'
sheet.cell(row=1,column=5).value = 'End long'

sheet.cell(row=1,column=6).value = 'GG start point distance'
sheet.cell(row=1,column=7).value = 'GG end point distance'

sheet.cell(row=1,column=8).value = 'GHTK car start point distance'
sheet.cell(row=1,column=9).value = 'GGHTK car end point distance'

sheet.cell(row=1,column=10).value = 'GHTK bike start point distance'
sheet.cell(row=1,column=11).value = 'GHTK bike end point distance'

sheet.cell(row=1,column=12).value = 'GHTK motor start point distance'
sheet.cell(row=1,column=13).value = 'GHTK motor end point distance'

sheet.cell(row=1,column=14).value = 'GHTK xteam motor start point distance'
sheet.cell(row=1,column=15).value = 'GHTK xteam motor end point distance'

sheet.cell(row=1,column=16).value = 'GH car error'
sheet.cell(row=1,column=17).value = 'GH bike error'
sheet.cell(row=1,column=18).value = 'GH motor error'
sheet.cell(row=1,column=19).value = 'GH xteam motor error'

for i in range(2,sheet_gg.max_row+1):
    print("Case: "+str(sheet_gg.cell(row=i,column=1).value))

    start = (sheet_gg.cell(row=i,column=2).value,sheet_gg.cell(row=i,column=3).value) # long, lat
    end = (sheet_gg.cell(row=i,column=4).value,sheet_gg.cell(row=i,column=5).value) # long, lat
    
    sheet.cell(row=i,column=1).value = sheet_gg.cell(row=i,column=1).value
    sheet.cell(row=i,column=2).value = sheet_gg.cell(row=i,column=3).value
    sheet.cell(row=i,column=3).value = sheet_gg.cell(row=i,column=2).value
    sheet.cell(row=i,column=4).value = sheet_gg.cell(row=i,column=5).value
    sheet.cell(row=i,column=5).value = sheet_gg.cell(row=i,column=4).value
    
    num_points = 500  # Number of points to interpolate
    
    # google_line = wkt.loads(sheet_gg.cell(row=i,column=7).value)
    
    google_encoded = str(sheet_gg.cell(row=i,column=6).value)
    google_decoded = polyline.decode(google_encoded)
    google_line = LineString(google_decoded)
    
    google_line = LineString([[lon, lat] for lat, lon in google_line.coords])
    
    gg_start = google_line.coords[0]
    gg_end = google_line.coords[-1]
    
    sheet.cell(row=i,column=6).value = haversine_for_points(start,gg_start) * 1000
    sheet.cell(row=i,column=7).value = haversine_for_points(end,gg_end) * 1000
    
    # GraphHopper API accepts: car, bike, foot
    gh_response_car = ghtk_gh_api(start,end,'car')
    # time.sleep(20)
    gh_response_bike = ghtk_gh_api(start,end,'bike')
    # time.sleep(20)
    gh_response_foot = ghtk_gh_api(start,end,'motorcycle')
    # time.sleep(20)
    gh_response_motor = ghtk_gh_api(start,end,'xteam_motorcycle')
    # time.sleep(20)
    
    if gh_response_car['status'] is False:
        gh_line_car = LineString([])
        error_message_car = gh_response_car['message']
    else:
        gh_line_car = LineString(gh_response_car['route'])
        error_message_car = gh_response_car['message']
    
    if gh_response_bike['status'] is False:
        gh_line_bike = LineString([])
        error_message_bike = gh_response_bike['message']
    else:
        gh_line_bike = LineString(gh_response_bike['route'])
        error_message_bike = gh_response_bike['message']
    
    if gh_response_foot['status'] is False:
        gh_line_foot = LineString([])
        error_message_foot = gh_response_foot['message']
    else:
        gh_line_foot = LineString(gh_response_foot['route'])
        error_message_foot = gh_response_foot['message']
        
    if gh_response_motor['status'] is False:
        gh_line_motor = LineString([])
        error_message_motor = gh_response_motor['message']
    else:
        gh_line_motor = LineString(gh_response_motor['route'])
        error_message_motor = gh_response_motor['message']
    
    geometry = [gh_line_car,gh_line_bike,gh_line_foot,gh_line_motor]
    error = [error_message_car,error_message_bike,error_message_foot,error_message_motor]
    
    for x in range(len(geometry)):
        if LineString(geometry[x]).is_empty:
            sheet.cell(row=i,column=16+x).value = error[x]
        else:
            sheet.cell(row=i,column=16+x).value = error[x]
            
            gh_start = geometry[x].coords[0]
            gh_end = geometry[x].coords[-1]
            
            sheet.cell(row=i,column=8+(2*x)).value = haversine_for_points(start,gh_start) * 1000
            sheet.cell(row=i,column=9+(2*x)).value = haversine_for_points(end,gh_end) * 1000

sheet_file = f"C:\\Users\\phams\\Downloads\\linestring_analysis_start_end_{today}.xlsx"
spreadsheet.save(sheet_file)