import requests
from shapely.geometry import LineString, Point
from shapely.ops import split
from shapely import wkt
import numpy as np
import json
import polyline
import ast
import matplotlib.pyplot as plt
import contextily as ctx
import geopandas as gpd
import psycopg2
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from openpyxl import Workbook
from openpyxl.drawing.image import Image
import io
import os
import pandas as pd
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from PIL import Image
from pyproj import Proj, transform
from geojson_length import calculate_distance, Unit
from geojson import Feature, LineString as GeoJSONLineString
import openpyxl
from scipy.spatial.distance import directed_hausdorff
from docx.shared import Pt



# Google excel sheet
spreadsheet_gg = openpyxl.load_workbook("C:\\Users\\phams\\Downloads\\failed_case.xlsx")

sheet_gg = spreadsheet_gg.active

# document
doc = Document()

# sheet
spreadsheet = openpyxl.Workbook()
sheet = spreadsheet.active

img_stream = io.BytesIO()



def ghtk_gh_api(start,end,type):
    headers = {
        'apikey': 'aghzzzzzzzzzzzzui',
        'Content-Type': 'application/json',
        'Authorization': 'Bearer eyJhbGciOiJFUzI1NiIsImtpZCI6IjAxRjRTMFBTRjdQR001VDZDWVExTTVNRVg3XzE2MjAwNDIyNzgiLCJ0eXAiOiJKV1QifQ.eyJhdWQiOiJnbWFwIiwiZXhwIjoxNjcwMzg4OTg0LCJqdGkiOiIwMUdLSldOS1NTN0FDSldTTVExMURHNzhZRSIsImlhdCI6MTY3MDMwMjU3OSwiaXNzIjoiaHR0cHM6Ly9hdXRoLmdodGtsYWIuY29tIiwic3ViIjoiMDFGNFMwUFNGN1BHTTVUNkNZUTFNNU1FWDciLCJzY3AiOlsiZ21hcDphZGRyZXNzLXZlcmlmaWNhdGlvbi1hcGkuYWNjZXNzLWRpcmVjdCIsImdtYXA6YWRkcmVzc2VzLW1hbmFnZW1lbnQuYWNjZXNzLWRpcmVjdCIsImdtYXA6Y2FwYWNpdHkuYWNjZXNzLWRpcmVjdCIsImdtYXA6Z29vZ2xlLW1hcC1odHRwLXNlcnZpY2UuYWNjZXNzLWRpcmVjdCIsImdtYXA6aGFtbGV0LW5vZGUtZGV0YWlsLmFjY2Vzcy1kaXJlY3QiLCJnbWFwOmludGVybmFsLWFwaS5hY2Nlc3MtZGlyZWN0IiwiZ21hcDpwcmVnZW8uYWNjZXNzLWRpcmVjdCIsImdtYXA6cm91dGUuYWNjZXNzLWRpcmVjdCIsImdtYXA6c3VnZ2VzdC1hZGRyZXNzLmFjY2Vzcy1kaXJlY3QiXSwiY2xpZW50X2lkIjoiMDFGNFMwUFNGN1BHTTVUNkNZUTFNNU1FWDciLCJ0eXBlIjoiZGlyZWN0In0.b89nX-mK6pjtfAMtqcCXFghvPC1qbOHjPDaDDrz9ljuTpdAhkvxkVlnST2tSyxNJ4dluwjHIBPsmFzWS5erJUA',
    }

    json_data = {
        'gh_requests': [
            {
                'points': [
                    [
                        start[1],start[0]
                    ],
                    [
                        end[1],end[0]
                    ],
                ],
                'profile': type,
                'request_id': '982884285',
                'calc_points': True,
                'points_encoded': True,
                'instructions': True,
            },
        ],
    }

    response = requests.post('http://10.110.69.186:8989/route', headers=headers, json=json_data)

    # Note: json_data will not be serialized by requests
    # exactly as it was in the original request.
    #data = '{\n        "gh_requests": [\n            {\n                "points": [\n                    [\n                        105.842638,\n                        21.200526\n                    ],\n                    [\n                        104.1166,\n                        21.185161\n                    ]\n                ],\n                "profile": "car",\n                "request_id": "982884285",\n                "calc_points": true,\n                "points_encoded": true,\n                "instructions": true\n            }\n        ]\n    }'
    #response = requests.post('http://10.110.69.186:8989/route', headers=headers, data=data)

    data = response.json()
    
    if data.get('gh_responses') is None:
        data = {
            'status': False,
            'route': '',
            'message': "Error getting response"
        }
        
        return data
    else:
        gh_response = data['gh_responses']
        if gh_response[0].get('paths', None) is None or gh_response[0]['paths'] == []:
            data = {
                'status': False,
                'route': '',
                'message': gh_response[0]['errors'][0]['message']
            }
            
            return data
        else:
            paths_data = gh_response[0]['paths']
            encoded_points = paths_data[0]['points']
            decoded_points = polyline.decode(encoded_points)
            
            data = {
                'status': True,
                'route': [(lon, lat) for lat, lon in decoded_points],
                'message': "Success"
            }
            
            return data

# Define a function to simplify a LineString
def simplify_route(line, tolerance=0.001):
    return line.simplify(tolerance, preserve_topology=False)

# Define a function to compute the Hausdorff distance
def hausdorff_distance(line1, line2):
    u = np.array(line1.coords)
    v = np.array(line2.coords)
    return max(directed_hausdorff(u, v)[0], directed_hausdorff(v, u)[0])

# Define a function to compute the average Hausdorff distance
def average_hausdorff_distance(line1, line2):
    u = np.array(line1.coords)
    v = np.array(line2.coords)
    distances_uv = [min(np.linalg.norm(u_i - v, axis=1)) for u_i in u]
    distances_vu = [min(np.linalg.norm(v_i - u, axis=1)) for v_i in v]
    avg_distance = (sum(distances_uv) + sum(distances_vu)) / (len(distances_uv) + len(distances_vu))
    return avg_distance

# Define a function to convert degrees to meters (for latitude and longitude)
def degrees_to_meters(degrees, latitude=0):
    meters_per_degree_lat = 111e3  # Approximate conversion: 1 degree of latitude = 111 km
    meters_per_degree_lon = meters_per_degree_lat * np.cos(np.radians(latitude))
    return degrees * meters_per_degree_lat, degrees * meters_per_degree_lon

def convert_to_lon_lat(linestring):
    return [[lon, lat] for lon, lat in linestring.coords]

def split_route_into_sections(line, percentage):
    total_length = line.length
    section_length = total_length * percentage / 100.0
    num_sections = int(np.ceil(total_length / section_length))
    points = [line.interpolate(float(n) / num_sections, normalized=True) for n in range(num_sections + 1)]
    sections = [(points[i], points[i+1]) for i in range(len(points) - 1)]
    return sections

def check_section_overlap(route_section, route):
    section_line = LineString(route_section)
    return section_line.intersects(route)

def check_route_through_sections(route_1, route_2, percentage):
    sections_1 = split_route_into_sections(route_1, percentage)
    
    for section in sections_1:
        if check_section_overlap(section, route_2):
            return True
    
    return False



doc.add_heading('LineString Analysis', level=1)

sheet.cell(row=1,column=1).value = 'Case'

sheet.cell(row=1,column=2).value = 'Start lat'
sheet.cell(row=1,column=3).value = 'Start long'

sheet.cell(row=1,column=4).value = 'End lat'
sheet.cell(row=1,column=5).value = 'End long'

sheet.cell(row=1,column=6).value = 'GHTK car error'
sheet.cell(row=1,column=7).value = 'GHTK bike error'
sheet.cell(row=1,column=8).value = 'GHTK motorcycle error'
sheet.cell(row=1,column=9).value = 'GHTK xteam motorcycle error'

for i in range(2,sheet_gg.max_row+1):
    print("Case: "+str(sheet_gg.cell(row=i,column=1).value))
    # doc.add_heading(f"Case: {str(sheet_gg.cell(row=i,column=1).value)}", level=2)
    sheet.cell(row=i,column=1).value = sheet_gg.cell(row=i,column=1).value

    start = (sheet_gg.cell(row=i,column=3).value,sheet_gg.cell(row=i,column=2).value) # lat, long
    end = (sheet_gg.cell(row=i,column=5).value,sheet_gg.cell(row=i,column=4).value) # lat, long
    
    # doc.add_paragraph(f"Start at: lat = {sheet_gg.cell(row=i,column=3).value}, long = {sheet_gg.cell(row=i,column=2).value}")
    # doc.add_paragraph(f"End at: lat = {sheet_gg.cell(row=i,column=5).value}, long = {sheet_gg.cell(row=i,column=4).value}")
    sheet.cell(row=i,column=2).value = sheet_gg.cell(row=i,column=3).value
    sheet.cell(row=i,column=3).value = sheet_gg.cell(row=i,column=2).value
    sheet.cell(row=i,column=4).value = sheet_gg.cell(row=i,column=5).value
    sheet.cell(row=i,column=5).value = sheet_gg.cell(row=i,column=4).value

    # GHTK API accepts: car, bike, motorcycle, xteam_motorcycle
    ghtk_response_car = ghtk_gh_api(start,end,'car')
    ghtk_response_bike = ghtk_gh_api(start,end,'bike')
    ghtk_response_motorcycle = ghtk_gh_api(start,end,'motorcycle')
    ghtk_response_xteam_motorcycle = ghtk_gh_api(start,end,'xteam_motorcycle')
    
    if ghtk_response_car['status'] is False:
        ghtk_line_car = LineString([])
        error_message_car = ghtk_response_car['message']
    else:
        ghtk_line_car = LineString(ghtk_response_car['route'])
        error_message_car = "No error"
    
    if ghtk_response_bike['status'] is False:
        ghtk_line_bike = LineString([])
        error_message_bike = ghtk_response_bike['message']
    else:
        ghtk_line_bike = LineString(ghtk_response_bike['route'])
        error_message_bike = "No error"
    
    if ghtk_response_motorcycle['status'] is False:
        ghtk_line_motorcycle = LineString([])
        error_message_motorcycle = ghtk_response_motorcycle['message']
    else:
        ghtk_line_motorcycle = LineString(ghtk_response_motorcycle['route'])
        error_message_motorcycle = "No error"
    
    if ghtk_response_xteam_motorcycle['status'] is False:
        ghtk_line_xteam_motorcycle = LineString([])
        error_message_xteam_motorcycle = ghtk_response_xteam_motorcycle['message']
    else:
        ghtk_line_xteam_motorcycle = LineString(ghtk_response_xteam_motorcycle['route'])
        error_message_xteam_motorcycle = "No error"
    
    google_line = wkt.loads(sheet_gg.cell(row=i,column=7).value)
    
    ls_data = {
        'geometry': [google_line,ghtk_line_car,ghtk_line_bike,ghtk_line_motorcycle,ghtk_line_xteam_motorcycle],
        'route_name': ['Google','GHTK car','GHTK bike','GHTK motorcycle','GHTK xteam motorcycle'],
        'error_message': ['No error',error_message_car,error_message_bike,error_message_motorcycle,error_message_xteam_motorcycle]
    }
    
    gdf = gpd.GeoDataFrame(data=ls_data, crs="EPSG:4326")

    # doc.add_heading('Route Lengths', level=3)

    for x in range(1,len(gdf)):
        sheet.cell(row=i,column=5+x).value = gdf.iloc[x].error_message
        # ls = gdf.iloc[x].geometry
        # name = gdf.iloc[x].route_name
        # error = gdf.iloc[x].error_message
        
        # if not LineString(ls).is_empty:
        #     ls_converted = convert_to_lon_lat(ls)
        #     line = Feature(geometry=GeoJSONLineString(ls_converted))
        #     length = calculate_distance(line,Unit.kilometers)
            
        #     print(f"Length of {name} route: {length:.2f} kilometer")
            
        #     doc.add_paragraph(f"Length of {name} route: {length:.2f} kilometer")
        #     sheet.cell(row=i,column=6+x).value = length
        # else:
        #     print(f"{name} didn't return a value")
            
        #     doc.add_paragraph(f"{name} didn't return a value: {error}")

    # doc.add_heading('Average Distances', level=3)
    
    # for y in range(1,len(gdf)):
    #     ls1 = gdf.iloc[y].geometry
    #     name1 = gdf.iloc[y].route_name
        
    #     doc.add_heading(f"Hausdorff distance between {name1} and Google", level=4)
        
    #     if LineString(ls1).is_empty:
    #         print(f"{name1} didn't return a value, therefore Hausdorff distance to Google route can't be calculated")
            
    #         doc.add_paragraph(f"{name1} didn't return a value, therefore Hausdorff distance to Google route can't be calculated")
    #     else:
    #         simplified_ls_gg = simplify_route(google_line)
    #         simplified_ls1 = simplify_route(ls1)
            
    #         # Calculate Hausdorff distance in degrees
    #         distance_degrees = hausdorff_distance(simplified_ls1, simplified_ls_gg)

    #         # Calculate average Hausdorff distance in degrees
    #         avg_distance_degrees = average_hausdorff_distance(simplified_ls1, simplified_ls_gg)
            
    #         latitude = (google_line.centroid.y + ls1.centroid.y) / 2
    #         distance_meters_lat, distance_meters_lon = degrees_to_meters(distance_degrees, latitude)
    #         avg_distance_meters_lat, avg_distance_meters_lon = degrees_to_meters(avg_distance_degrees, latitude)
                    
    #         print(f"Hausdorff distance between {name1} and Google route in degree: {distance_degrees}")
    #         doc.add_paragraph(f"Hausdorff distance between {name1} and Google route in degree: {distance_degrees}")
    #         sheet.cell(row=i,column=10+y).value = distance_degrees
            
    #         print(f"Hausdorff distance between {name1} and Google route in meters (latitude): {distance_meters_lat}")
    #         doc.add_paragraph(f"Hausdorff distance between {name1} and Google route in meters (latitude): {distance_meters_lat}")
            
    #         print(f"Hausdorff distance between {name1} and Google route in meters (longitude): {distance_meters_lon}")
    #         doc.add_paragraph(f"Hausdorff distance between {name1} and Google route in meters (longitude): {distance_meters_lon}")
            
    #         print(f"Average Hausdorff distance between {name1} and Google route in degree: {avg_distance_degrees}")
    #         doc.add_paragraph(f"Average Hausdorff distance between {name1} and Google route in degree: {avg_distance_degrees}")
    #         sheet.cell(row=i,column=14+y).value = avg_distance_degrees
            
    #         print(f"Average Hausdorff distance between {name1} and Google route in meters (latitude): {avg_distance_meters_lat}")
    #         doc.add_paragraph(f"Average Hausdorff distance between {name1} and Google route in meters (latitude): {avg_distance_meters_lat}")
            
    #         print(f"Average Hausdorff distance between {name1} and Google route in meters (longitude): {avg_distance_meters_lon}")
    #         doc.add_paragraph(f"Average Hausdorff distance between {name1} and Google route in meters (longitude): {avg_distance_meters_lon}")
            
            # overlap = check_route_through_sections(ls1,google_line,5.0)
              
            # if overlap:
            #     print(f"The {name1} route goes through one or more sections of the Google route route")
                        
            #     doc.add_paragraph(f"The {name1} route goes through one or more sections of the Google route route")
            # else:
            #     print(f"The {name1} route does not go through any sections of the Google route route")
                        
            #     doc.add_paragraph(f"The {name1} route does not go through any sections of the Google route route")
    
    # colors = ['red','blue','green','purple','orange']

    # fig, ax = plt.subplots()

    # for idx, row in gdf.iterrows():
    #     # Plot the LineString
    #     if not row['geometry'].is_empty:
    #         gdf.loc[[idx]].plot(ax=ax, label=row['route_name'], color=colors[idx % len(colors)], linewidth=2)

    # # Set a fixed aspect ratio
    # ax.set_aspect('equal')

    # # Customize and show the plot
    # plt.title('Route from Point A to Point B')
    # plt.xlabel('Longitude')
    # plt.ylabel('Latitude')
    # plt.legend(loc='lower right')
    # ctx.add_basemap(ax, crs="EPSG:4326", source=ctx.providers.OpenStreetMap.Mapnik, zoom=15)
    # plt.grid(True)

    # # manager = plt.get_current_fig_manager()
    # # manager.window.state('zoomed')

    # # plt.show()

    # plt.savefig(img_stream, format='png')
    # # Save the plot to a file
    # plot_file = "C:\\Users\\phams\\Downloads\\linestrings_plot.png"
    # fig.savefig(plot_file)
    # plt.close(fig)  # Close the plot

    # Check if the file is created and valid
    # try:
    #     with Image.open(plot_file) as img:
    #         img.show()  # Open the image to verify
    # except IOError:
    #     print("Error: The image was not saved correctly.")

    # doc.add_heading('Plot of LineStrings', level=2)
    # doc.add_picture(plot_file, width=Inches(5))
    
    # doc.add_page_break()
    
# # Define the desired font size
# font_size = Pt(12)  # 12 point font size
# # Set the font size for all paragraphs and runs in the document
# for paragraph in doc.paragraphs:
#     for run in paragraph.runs:
#         run.font.size = font_size
# doc_file = "C:\\Users\\phams\\Downloads\\linestring_analysis.docx" # for text only
# # doc_file = "C:\\Users\\phams\\Downloads\\linestring_analysis_graphs.docx" # for graph only
# doc.save(doc_file)

sheet_file = "C:\\Users\\phams\\Downloads\\linestring_analysis_error.xlsx"
spreadsheet.save(sheet_file)

